#!/usr/bin/env python3
"""
Word 报告生成模块
- 左右分栏双语对照（英文 / 中文翻译）
- 6维度深度分析表格
- 自动完整性检查与修复
"""

import argparse
import json
import logging
import os
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# ── 依赖检查 ──────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word 报告生成不可用。请运行: pip install python-docx")


# ── 样式工具 ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """设置单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, border_color: str = "CCCCCC", border_size: int = 4):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{edge}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), str(border_size))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), border_color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _heading_paragraph(doc, text: str, level: int = 1, color: str = "003366"):
    """添加带颜色的标题段落。"""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_cover(doc, topic_name: str, date_range: str, paper_count: int):
    """添加封面页。"""
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{topic_name} 文献调研报告")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # 副信息
    for label, value in [
        ("时间范围", date_range),
        ("文献数量", f"{paper_count} 篇"),
        ("数据来源", "PubMed"),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()


def _add_paper_section(doc, paper: dict, idx: int):
    """添加单篇文献的完整分析节。"""
    title = paper.get("title", "")
    title_cn = paper.get("title_cn", "")
    abstract = paper.get("abstract", "")
    abstract_cn = paper.get("abstract_cn", "")
    journal = paper.get("journal", "")
    if_val = paper.get("journal_if", "—")
    pub_date = paper.get("publication_date", "—")
    author_display = paper.get("author_display", "")
    doi = paper.get("doi", "")
    pmid = paper.get("pmid", "")
    # ── 标题 ──────────────────────────────────────────────────────────────────
    _heading_paragraph(doc, f"{idx}. {title}", level=2)
    if title_cn:
        p = doc.add_paragraph()
        p.add_run("中文标题：").font.bold = True
        p.add_run(title_cn)

    # ── 基本信息表 ────────────────────────────────────────────────────────────
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"

    def _add_info_row(table, label: str, value: str):
        row = table.add_row()
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        _set_cell_bg(label_cell, "EEF2FF")
        label_cell.width = Cm(3.5)
        value_cell.width = Cm(13)
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.bold = True
        lr.font.size = Pt(9)
        vp = value_cell.paragraphs[0]
        vp.add_run(value).font.size = Pt(9)

    _add_info_row(info_table, "期刊", f"{journal}（IF: {if_val}）")
    _add_info_row(info_table, "发表日期", pub_date)
    _add_info_row(info_table, "作者", author_display)
    if doi:
        _add_info_row(info_table, "DOI", doi)
        _add_info_row(info_table, "原文链接", f"https://doi.org/{doi}")
    if pmid:
        _add_info_row(info_table, "PubMed", f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")

    doc.add_paragraph()

    # ── 摘要双栏对照 ──────────────────────────────────────────────────────────
    _heading_paragraph(doc, "摘要对照", level=3, color="0055AA")

    if abstract or abstract_cn:
        abs_table = doc.add_table(rows=1, cols=2)
        abs_table.style = "Table Grid"
        abs_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        hdr = abs_table.rows[0]
        for cell, text in zip(hdr.cells, ["英文原文", "中文翻译"]):
            _set_cell_bg(cell, "003366")
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 内容行
        content_row = abs_table.add_row()
        left_cell, right_cell = content_row.cells
        left_cell.width = Cm(8)
        right_cell.width = Cm(8)

        lp = left_cell.paragraphs[0]
        lp.add_run(abstract or "（无摘要）").font.size = Pt(9)
        rp = right_cell.paragraphs[0]
        rp.add_run(abstract_cn or "（无翻译）").font.size = Pt(9)
        doc.add_paragraph()

    # ── 6维度分析 ─────────────────────────────────────────────────────────────
    dims = [
        ("technical_route", "技术路线"),
        ("advantages", "技术优势"),
        ("limitations", "技术不足"),
        ("technical_barriers", "技术壁垒"),
        ("feasibility", "落地可行性"),
        ("generalization", "泛化能力"),
    ]
    has_analysis = any(paper.get(k) for k, _ in dims)
    if has_analysis:
        _heading_paragraph(doc, "深度分析（6维度）", level=3, color="0055AA")
        dim_table = doc.add_table(rows=0, cols=2)
        dim_table.style = "Table Grid"

        for key, label in dims:
            content = paper.get(key, "").strip() or "待分析"
            row = dim_table.add_row()
            label_cell, content_cell = row.cells
            _set_cell_bg(label_cell, "E8F4FD")
            label_cell.width = Cm(3)
            content_cell.width = Cm(13.5)
            lp = label_cell.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = RGBColor(0, 51, 102)
            cp = content_cell.paragraphs[0]
            cp.add_run(content).font.size = Pt(9)

    doc.add_paragraph()

    # 分隔线（通过段落样式模拟）
    p = doc.add_paragraph("─" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(8)
    doc.add_paragraph()


# ── 完整性检查与修复 ──────────────────────────────────────────────────────────

def _check_paper_completeness(paper: dict) -> list[str]:
    """返回该文献存在的完整性问题列表。"""
    issues = []
    if not paper.get("title") or len(paper.get("title", "")) < 10:
        issues.append("标题缺失或过短")
    if not paper.get("abstract") or len(paper.get("abstract", "")) < 100:
        issues.append(f"摘要缺失或过短（{len(paper.get('abstract', ''))}字符）")
    if not paper.get("abstract_cn"):
        issues.append("中文翻译缺失")
    for key, label in [
        ("technical_route", "技术路线"),
        ("advantages", "技术优势"),
        ("limitations", "技术不足"),
        ("technical_barriers", "技术壁垒"),
        ("feasibility", "落地可行性"),
        ("generalization", "泛化能力"),
    ]:
        content = paper.get(key, "")
        if not content or len(content) < 20 or content in ("分析失败", "该维度分析内容待补充", "摘要缺失，无法分析"):
            issues.append(f"{label}分析不完整")
    return issues


def _auto_fix_paper(paper: dict) -> dict:
    """尝试自动修复文献的不完整字段。"""
    script_dir = Path(__file__).parent
    sys.path.insert(0, str(script_dir))

    from utils import translate_text
    from analyze_content import analyze_paper_content, validate_analysis_complete
    from fetch_papers import _get_abstract_from_europepmc

    issues = _check_paper_completeness(paper)
    if not issues:
        return paper

    logger.info("自动修复: %s", paper.get("title", "")[:50])

    # 修复摘要
    if any("摘要" in i for i in issues):
        doi = paper.get("doi")
        if doi:
            abstract = _get_abstract_from_europepmc(doi)
            if abstract and len(abstract) > 100:
                paper["abstract"] = abstract
                logger.info("  ✅ 重新获取摘要: %d 字符", len(abstract))

    # 修复翻译
    if any("翻译" in i for i in issues) and paper.get("abstract"):
        paper["abstract_cn"] = translate_text(paper["abstract"], type="abstract")
        logger.info("  ✅ 重新翻译摘要")

    # 修复6维度分析
    dim_keys = {"technical_route", "advantages", "limitations", "technical_barriers", "feasibility", "generalization"}
    needs_analysis = any(
        paper.get(k, "")
        in ("", "分析失败", "该维度分析内容待补充", "摘要缺失，无法分析")
        or len(paper.get(k, "")) < 20
        for k in dim_keys
    )
    if needs_analysis and paper.get("abstract"):
        time.sleep(2)
        analysis = analyze_paper_content(paper["title"], paper["abstract"], paper.get("journal", ""))
        if validate_analysis_complete(analysis):
            paper.update(analysis)
            logger.info("  ✅ 重新完成6维度分析")

    return paper


# ── 公开 API ──────────────────────────────────────────────────────────────────

def create_word_report(
    papers: list[dict],
    output_path: str,
    topic_name: str,
    date_range: str,
    auto_fix: bool = True,
) -> Tuple[bool, str]:
    """
    生成 Word 文档报告。

    Args:
        papers: 文献列表
        output_path: .docx 输出路径
        topic_name: 主题中文名称
        date_range: 时间范围字符串
        auto_fix: 是否自动修复不完整内容

    Returns:
        (success, message)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx 未安装"

    if not papers:
        return False, "无文献数据"

    # 完整性检查 & 修复
    if auto_fix:
        issues_count = sum(1 for p in papers if _check_paper_completeness(p))
        if issues_count:
            logger.info("发现 %d/%d 篇文献内容不完整，开始自动修复...", issues_count, len(papers))
            papers = [_auto_fix_paper(p) for p in papers]

    try:
        doc = Document()

        # 页面设置：A4
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # 封面
        _add_cover(doc, topic_name, date_range, len(papers))

        # 文献总览
        _heading_paragraph(doc, "文献概览", level=1)
        overview_table = doc.add_table(rows=1, cols=5)
        overview_table.style = "Table Grid"
        headers = ["序号", "标题（中英）", "期刊 / IF", "发表日期", "DOI"]
        for i, h in enumerate(headers):
            cell = overview_table.rows[0].cells[i]
            _set_cell_bg(cell, "003366")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, paper in enumerate(papers, 1):
            row = overview_table.add_row()
            title = paper.get("title", "")[:60]
            title_cn = paper.get("title_cn", "")[:40]
            title_display = f"{title}\n{title_cn}" if title_cn else title
            journal = paper.get("journal", "")
            if_val = paper.get("journal_if", "—")
            pub_date = paper.get("publication_date", "—")
            doi = paper.get("doi", "—")
            values = [str(idx), title_display, f"{journal}\n(IF: {if_val})", pub_date, doi]
            for i, val in enumerate(values):
                cp = row.cells[i].paragraphs[0]
                cp.add_run(val).font.size = Pt(8)
                if idx % 2 == 0:
                    _set_cell_bg(row.cells[i], "F5F8FF")

        doc.add_page_break()

        # 各篇分析
        _heading_paragraph(doc, "单篇文献详细分析", level=1)
        for idx, paper in enumerate(papers, 1):
            _add_paper_section(doc, paper, idx)

        # 保存
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return True, f"Word 报告已生成: {output_path}"

    except Exception as e:
        logger.exception("Word 报告生成失败")
        return False, str(e)


# ── CLI ───────────────────────────────────────────────────────────────────────

def _main():
    parser = argparse.ArgumentParser(description="生成文献调研 Word 报告")
    parser.add_argument("--input", required=True, help="papers JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--topic-name", default="未命名主题", help="主题名称")
    parser.add_argument("--date-range", default="", help="时间范围")
    parser.add_argument("--no-fix", action="store_true", help="禁用自动修复")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(message)s")

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)
    papers = data if isinstance(data, list) else data.get("papers", [])

    success, msg = create_word_report(
        papers, args.output, args.topic_name,
        args.date_range or "—", auto_fix=not args.no_fix,
    )
    print(("✅ " if success else "❌ ") + msg)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    _main()
